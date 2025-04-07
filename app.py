from flask import Flask, request, send_file, render_template_string
from docx import Document
import tempfile
import os

app = Flask(__name__)
UPLOAD_FOLDER = tempfile.gettempdir()

HTML_FORM = """
<!doctype html>
<title>DOCX Placeholder Replacer</title>
<h2>Upload DOCX and Provide Replacements</h2>
<form method=post enctype=multipart/form-data>
  <input type=file name=docx_file required><br><br>
  <textarea name=replacements placeholder='{"name": "Alice", "date": "April 7, 2025"}' rows=10 cols=50 required></textarea><br><br>
  <input type=submit value=Replace>
</form>
"""

def replace_placeholders(doc_path, output_path, replacements):
    doc = Document(doc_path)
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if f"{{{{{key}}}}}" in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if f"{{{{{key}}}}}" in inline[i].text:
                        inline[i].text = inline[i].text.replace(f"{{{{{key}}}}}", val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", val)

    doc.save(output_path)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        file = request.files['docx_file']
        replacements = request.form['replacements']
        try:
            replacements = eval(replacements)
        except Exception as e:
            return f"Invalid JSON: {str(e)}"

        if file.filename.endswith('.docx'):
            input_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(input_path)
            output_path = os.path.join(UPLOAD_FOLDER, f"output_{file.filename}")
            replace_placeholders(input_path, output_path, replacements)
            return send_file(output_path, as_attachment=True)
        else:
            return "Please upload a .docx file."
    return render_template_string(HTML_FORM)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
